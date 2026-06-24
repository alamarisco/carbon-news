#!/usr/bin/env python3
"""Render the weekly 每週國際新聞蒐集 .docx in the ESTC house style.

Input: a JSON file describing the week's stories.
    {
      "year": 2026,
      "week_label": "2026.06.15-06.18",     # used in filename
      "update_date": "2026/6/18",           # shown on the 更新 line
      "out_dir": "./out",   # optional; default "./out". Override with --out-dir.
      "stories": [
        {
          "headline": "歐盟理事會採取措施加強歐盟碳邊境調節機制",
          "paragraphs": [
            {"runs": [
                {"t": "歐盟理事會今天就加強碳邊境調節機制(CBAM)達成一致立場。"},
                {"t": "CBAM 適用範圍將涵蓋更多新產品。", "bold": true}
            ]}
          ],
          "source_url": "https://...",
          "date": "2026/06/12"
        }
      ]
    }

Bold key points by setting "bold": true on the relevant run.
Output: <out_dir>/每週國際新聞蒐集_<week_label>.docx
"""
import sys, json, os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

FONT = "Times New Roman"


def set_font(run, name=FONT):
    run.font.name = name
    rpr = run._element.rPr
    if rpr is not None:
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), name)


def add_line(doc, segments, size, default_bold=False):
    """segments: list of (text, bold) ; size in pt."""
    p = doc.add_paragraph()
    for text, bold in segments:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bool(bold) if bold is not None else default_bold
        set_font(r)
    return p


def main(spec_path, out_dir_override=None):
    spec = json.load(open(spec_path, encoding="utf-8"))
    doc = Document()

    # Header block
    add_line(doc, [("產品溫室氣體排放強度建立及碳邊境調整機制推動計畫", True)], 18)
    add_line(doc, [("因應計畫內容蒐集國際間最新推動資訊定期更新", True)], 18)
    add_line(doc, [(f"{spec['update_date'].replace('/', '年', 1).replace('/', '月', 1)}日更新", False)], 14)
    n = len(spec["stories"])
    add_line(doc, [("本週共計 ", False), (str(n), False), (" ", False), ("則", True)], 16)
    doc.add_paragraph()

    for s in spec["stories"]:
        # Headline — try named style, fall back to bold 16pt
        hp = doc.add_paragraph()
        try:
            hp.style = doc.styles["新聞標題"]
            hr = hp.add_run(s["headline"]); set_font(hr)
        except KeyError:
            hr = hp.add_run(s["headline"]); hr.bold = True; hr.font.size = Pt(16); set_font(hr)

        for para in s["paragraphs"]:
            p = doc.add_paragraph()
            for run in para["runs"]:
                r = p.add_run(run["t"])
                r.font.size = Pt(14)
                r.bold = bool(run.get("bold", False))
                set_font(r)

        add_line(doc, [(f"新聞出處：{s['source_url']}", False)], 14)
        add_line(doc, [(f"日期：{s['date']}", False)], 14)
        doc.add_paragraph()

    out_dir = out_dir_override or spec.get("out_dir") or "./out"
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, f"每週國際新聞蒐集_{spec['week_label']}.docx")
    doc.save(out)
    print("Saved:", out)


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser(description="Render the weekly 國際新聞蒐集 .docx")
    ap.add_argument("spec", help="spec.json describing the week's stories")
    ap.add_argument("--out-dir", default=None,
                    help="output dir (overrides spec.out_dir; default ./out)")
    a = ap.parse_args()
    main(a.spec, a.out_dir)
