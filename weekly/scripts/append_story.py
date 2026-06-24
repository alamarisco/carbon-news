#!/usr/bin/env python3
"""
FLAG-mode helper: append one translated story to an existing weekly .docx in the
ESTC house style, and finalize the doc's formatting.

Usage:
    python append_story.py <in.docx> <content.txt> <out.docx>

content.txt markup (Traditional Chinese):
    # Headline          -> style 新聞標題, auto-numbered, bookmarked; also added to 項次 table
    ## Sub-header        -> Normal 14pt bold (use ONLY if the source article has section headers)
    SRC <url>           -> 新聞出處：<url>
    DATE <YYYY/MM/DD>    -> 日期：<...>
    <other line>        -> Normal 14pt body (plain, no bold, no first-line indent)

Also finalizes the whole document:
  * headlines keep the existing numbered-list format (numId reused);
  * body paragraphs have their first-line indent removed (house style = flush left);
  * the 項次 index table: each title becomes an internal hyperlink jumping to its article;
  * 本週共計 N 則 recalculated.
"""
import sys, re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

LATIN, EAST = "Times New Roman", "標楷體"
HEAD_LINE, BODY_LINE, BODY_BEFORE = 24, 23, 9
LINK_COLOR = "0563C1"


def set_font(run, size=None, bold=None):
    run.font.name = LATIN
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {}); rpr.append(rf)
    rf.set(qn("w:eastAsia"), EAST); rf.set(qn("w:ascii"), LATIN); rf.set(qn("w:hAnsi"), LATIN)


def fmt(p, line_pt, space_before_pt=None, first_line=None):
    pf = p.paragraph_format
    pf.line_spacing = Pt(line_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_after = None
    pf.space_before = Pt(space_before_pt) if space_before_pt is not None else None
    if first_line is not None:
        pf.first_line_indent = Pt(first_line)


def add_page_break(doc):
    p = doc.add_paragraph()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    p.add_run()._r.append(br)


def headline_numpr(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == "新聞標題":
            np = p._p.find(".//" + qn("w:numPr"))
            if np is not None:
                ilvl = np.find(qn("w:ilvl")); numid = np.find(qn("w:numId"))
                return (ilvl.get(qn("w:val")) if ilvl is not None else "0",
                        numid.get(qn("w:val")) if numid is not None else None)
    return None


def apply_numpr(p, ilvl, numid):
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    e1 = OxmlElement("w:ilvl"); e1.set(qn("w:val"), ilvl)
    e2 = OxmlElement("w:numId"); e2.set(qn("w:val"), numid)
    numPr.append(e1); numPr.append(e2); pPr.append(numPr)


def bookmark_para(p, name, bid):
    start = OxmlElement("w:bookmarkStart"); start.set(qn("w:id"), str(bid)); start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd"); end.set(qn("w:id"), str(bid))
    pPr = p._p.find(qn("w:pPr"))
    idx = list(p._p).index(pPr) + 1 if pPr is not None else 0
    p._p.insert(idx, start); p._p.append(end)


def cell_internal_link(cell, anchor, text, size):
    para = cell.paragraphs[0]
    for r in list(para._p.findall(qn("w:r"))):
        para._p.remove(r)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:eastAsia"), EAST); rf.set(qn("w:ascii"), LATIN); rf.set(qn("w:hAnsi"), LATIN); rPr.append(rf)
    col = OxmlElement("w:color"); col.set(qn("w:val"), LINK_COLOR); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if size:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; run.append(t)
    hl.append(run); para._p.append(hl)


def add_index_row(doc, num, headline):
    tbl = doc.tables[0]
    src = tbl.rows[-1]
    row = tbl.add_row()
    for ci, val in enumerate((str(num), headline)):
        scell = src.cells[ci]; spara = scell.paragraphs[0]
        srun = spara.runs[0] if spara.runs else None
        para = row.cells[ci].paragraphs[0]
        para.alignment = spara.alignment
        para.paragraph_format.line_spacing = spara.paragraph_format.line_spacing
        para.paragraph_format.line_spacing_rule = spara.paragraph_format.line_spacing_rule
        run = para.add_run(val)
        set_font(run, size=(srun.font.size.pt if srun and srun.font.size else None),
                 bold=(srun.bold if srun else None))


def normalize_body_indent(doc):
    for p in doc.paragraphs:
        if p.style and p.style.name == "Normal" and p.paragraph_format.first_line_indent is not None:
            p.paragraph_format.first_line_indent = Pt(0)


def linkify_index_table(doc):
    """Bookmark each headline in order; turn each table title cell into a jump link."""
    heads = [p for p in doc.paragraphs if p.style and p.style.name == "新聞標題" and p.text.strip()]
    for i, p in enumerate(heads, 1):
        bookmark_para(p, "art%d" % i, i)
    tbl = doc.tables[0]
    for ri in range(1, len(tbl.rows)):          # skip header row
        cell = tbl.rows[ri].cells[1]
        txt = cell.text.strip()
        sz = None
        if cell.paragraphs[0].runs and cell.paragraphs[0].runs[0].font.size:
            sz = cell.paragraphs[0].runs[0].font.size.pt
        if ri <= len(heads):
            cell_internal_link(cell, "art%d" % ri, txt, sz)


def main():
    inp, content, out = sys.argv[1], sys.argv[2], sys.argv[3]
    doc = Document(inp)
    numpr = headline_numpr(doc)
    headline = None

    add_page_break(doc)
    for raw in open(content, encoding="utf-8").read().splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            headline = line[2:].strip()
            p = doc.add_paragraph(); p.style = doc.styles["新聞標題"]
            set_font(p.add_run(headline)); fmt(p, HEAD_LINE)
            if numpr:
                apply_numpr(p, numpr[0], numpr[1])
        elif line.startswith("## "):
            p = doc.add_paragraph(); p.style = doc.styles["Normal"]
            set_font(p.add_run(line[3:].strip()), 14, True); fmt(p, BODY_LINE, BODY_BEFORE, 0)
        elif line.startswith("SRC "):
            p = doc.add_paragraph(); p.style = doc.styles["Normal"]
            set_font(p.add_run("新聞出處：" + line[4:].strip())); fmt(p, BODY_LINE, None, 0)
        elif line.startswith("DATE "):
            p = doc.add_paragraph(); p.style = doc.styles["Normal"]
            set_font(p.add_run("日期：" + line[5:].strip())); fmt(p, BODY_LINE, None, 0)
        else:
            p = doc.add_paragraph(); p.style = doc.styles["Normal"]
            set_font(p.add_run(line), 14, False); fmt(p, BODY_LINE, BODY_BEFORE, 0)

    add_index_row(doc, len(doc.tables[0].rows), headline)
    normalize_body_indent(doc)
    linkify_index_table(doc)

    n = sum(1 for p in doc.paragraphs if p.style and p.style.name == "新聞標題" and p.text.strip())
    for p in doc.paragraphs:
        if "本週共計" in p.text and "則" in p.text:
            for run in p.runs:
                run.text = re.sub(r"\d+", str(n), run.text)
            break

    doc.save(out)
    print("[OK] 項次%d appended; %d headlines bookmarked+linked; 本週共計 %d 則 -> %s"
          % (len(doc.tables[0].rows) - 1, n, n, out))


if __name__ == "__main__":
    main()
